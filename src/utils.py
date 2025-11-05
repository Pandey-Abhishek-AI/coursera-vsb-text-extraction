import os
import re
import requests
import tempfile
from io import BytesIO
from typing import Optional, Dict, Any, List, Tuple

import boto3
import markdown
import pandas as pd
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from bs4 import BeautifulSoup
from markitdown import MarkItDown

from dotenv import load_dotenv
load_dotenv()

from src.logging_config import get_logger
logger = get_logger(__name__)

# Initialize MarkItDown converter
md = MarkItDown()

S3_BUCKET = os.getenv("S3_BUCKET_NAME")
S3_REGION = os.getenv("AWS_REGION")

async def upload_file_to_s3(file_content: bytes, s3_key: str, content_type: str) -> str:
    """
    Utility function to upload a file to S3 and return the URL.
    """
    s3 = boto3.client("s3", region_name=S3_REGION)
    s3.put_object(
        Bucket=S3_BUCKET,
        Key=s3_key,
        Body=file_content,
        ContentType=content_type
    )
    public_url = f"https://{S3_BUCKET}.s3.{S3_REGION}.amazonaws.com/{s3_key}"
    return public_url

def convert_docx_to_markdown(file_path: str) -> Optional[str]:
    """
    Convert a DOCX file to markdown text.
    
    Args:
        file_path (str): Path to the DOCX file or BytesIO object
        
    Returns:
        Optional[str]: Converted markdown text, None if conversion fails
    """
    try:   
        # Convert the document to markdown
        markdown_text = md.convert(file_path)
        return markdown_text.text_content
        
    except Exception as e:
        logger.error(f"Error converting DOCX to markdown: {e}")
        return None


def convert_markdown_to_html(markdown_text: str, options: Optional[Dict[str, Any]] = None) -> str:
    """
    Convert Markdown text to HTML using the Python markdown library.
    
    Args:
        markdown_text (str): The markdown content to convert
        options (dict, optional): Configuration options for the markdown converter
        
    Returns:
        str: The HTML output
    """
    try:
        # Default options
        default_options = {
            'extensions': [
                'markdown.extensions.fenced_code',    # GitHub-style code blocks
                'markdown.extensions.tables',         # Tables support
                'markdown.extensions.nl2br',          # Convert newlines to <br>
                'markdown.extensions.sane_lists',     # Better list handling
                'markdown.extensions.smarty',         # Smart quotes, dashes, etc.
            ],
            'output_format': 'html',
            'border_style': 'css',
            'border_width': '1px',
            'border_color': '#dddddd',
            'border_style_value': 'solid'
        }
        
        # Merge default options with user-provided options
        if options is None:
            options = {}
        
        merged_options = {**default_options, **options}
        
        # Convert using the markdown library
        html = markdown.markdown(
            markdown_text, 
            extensions=merged_options['extensions'],
            output_format=merged_options['output_format']
        )

        # Add CSS styling for tables
        html = add_table_borders_with_css(html, merged_options)
        return html
    
    except Exception as e:
        logger.error(f"Error converting markdown to HTML: {e}")
        raise


def add_table_borders_with_css(html: str, options: Dict[str, Any]) -> str:
    """
    Add CSS styling for table borders to the HTML.
    
    Args:
        html (str): HTML content to process
        options (dict): Styling options
    
    Returns:
        str: HTML with added CSS for table borders
    """
    try:
        # Create a style tag with CSS for table borders
        border_width = options.get('border_width', '1px')
        border_color = options.get('border_color', '#dddddd')
        border_style = options.get('border_style_value', 'solid')
        
        css = f"""
    <style>
        table {{
            border-collapse: collapse;
            width: 100%;
            margin-bottom: 20px;
        }}
        th, td {{
            border: {border_width} {border_style} {border_color};
            padding: 8px;
            text-align: left;
        }}
        th {{
            background-color: #f2f2f2;
        }}
    </style>
    """
        
        # Add the CSS style tag to the HTML
        return css + html
        
    except Exception as e:
        logger.error(f"Error adding table borders: {e}")
        raise


def set_table_borders(table):
    """
    Add single-line borders on every side and between cells for a DOCX table.
    
    Args:
        table: DOCX table object
    """
    try:
        tbl = table._tbl
        tblPr = tbl.tblPr
        borders = OxmlElement("w:tblBorders")
        
        for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
            b = OxmlElement(f"w:{edge}")
            b.set(qn("w:val"), "single")
            b.set(qn("w:sz"), "4")       # 4 half-points = 2pt
            b.set(qn("w:color"), "000000")
            borders.append(b)
            
        tblPr.append(borders)
        
    except Exception as e:
        logger.error(f"Error setting table borders: {e}")
        raise


def html_table_to_docx(html_text: str, output_path: str) -> None:
    """
    Convert HTML tables to DOCX format with proper borders.
    
    Args:
        html_text (str): HTML content containing tables
        output_path (str): Path where the DOCX file will be saved
    """
    try:
        soup = BeautifulSoup(html_text, "html.parser")
        doc = Document()
        
        for tbl_html in soup.find_all("table"):
            rows = tbl_html.find_all("tr")
            if not rows:
                continue
                
            ncols = len(rows[0].find_all(["th", "td"]))
            table = doc.add_table(rows=0, cols=ncols)
            
            for row in rows:
                cells = row.find_all(["th", "td"])
                row_cells = table.add_row().cells
                for i, cell in enumerate(cells):
                    if i < len(row_cells):
                        row_cells[i].text = cell.get_text(strip=True)
                        
            set_table_borders(table)
            
        doc.save(output_path)
        logger.info(f"Successfully saved DOCX file: {output_path}")
        
    except Exception as e:
        logger.error(f"Error converting HTML table to DOCX: {e}")
        raise


def load_prompt(docx_path: str) -> Optional[str]:
    """
    Load the prompt text from a DOCX file.

    Args:
        docx_path (str): The path to the DOCX file

    Returns:
        Optional[str]: The prompt text as a string, None if loading fails
    """
    try:
        prompt_text = convert_docx_to_markdown(docx_path)
        
        return prompt_text
        
    except Exception as e:
        logger.error(f"Error loading prompt from {docx_path}: {e}")
        return None


def download_prompt_from_s3(s3_url: str) -> Optional[str]:
    """
    Download and load prompt text from an S3 URL.
    
    Args:
        s3_url (str): S3 URL of the prompt file
        
    Returns:
        Optional[str]: The prompt text as a string, None if loading fails
    """
    try:
        logger.info(f"Downloading prompt from S3: {s3_url}")
        
        # Download the file using requests (handles URL encoding automatically)
        response = requests.get(s3_url)
        response.raise_for_status()  # Raise an exception for bad status codes
        
        # Determine file type and process accordingly
        if s3_url.lower().endswith('.docx'):
            # Save to temporary file for DOCX processing
            with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as temp_file:
                temp_file.write(response.content)
                temp_file_path = temp_file.name
            
            try:
                prompt_text = convert_docx_to_markdown(temp_file_path)
                logger.info(f"Successfully loaded prompt from S3: {len(prompt_text) if prompt_text else 0} characters")
                return prompt_text
            finally:
                # Clean up temporary file
                os.unlink(temp_file_path)
                
        elif s3_url.lower().endswith(('.txt', '.md')):
            # Process text/markdown files directly
            prompt_text = response.content.decode('utf-8')
            logger.info(f"Successfully loaded prompt from S3: {len(prompt_text)} characters")
            return prompt_text
            
        else:
            logger.error(f"Unsupported file format for prompt: {s3_url}")
            return None
            
    except requests.exceptions.RequestException as e:
        logger.error(f"Error downloading prompt from S3 {s3_url}: {e}")
        return None
    except Exception as e:
        logger.error(f"Error processing prompt from S3 {s3_url}: {e}")
        return None


def slice_for_lesson(df: pd.DataFrame, lesson_num: str) -> pd.DataFrame:
    """
    Slice the DataFrame to extract data for a specific lesson.

    Args:
        df (pd.DataFrame): The input DataFrame containing the outline
        lesson_num (str): The lesson number to slice for

    Returns:
        pd.DataFrame: A DataFrame containing the rows for the specified lesson

    Raises:
        ValueError: If the header for the specified lesson is not found
    """
    try:
        tag = f"Lesson {lesson_num}"
        header_mask = df[0].astype(str).str.startswith(tag, na=False)
        
        if not header_mask.any():
            raise ValueError(f"Header '{tag}' not found in Column A.")
            
        start = header_mask.idxmax()

        # Find the next lesson header
        next_header = df[0].astype(str).str.match(r"Lesson\s+\d+", na=False)
        try:
            end = next_header[next_header & (next_header.index > start)].idxmax() - 1
        except ValueError:
            end = len(df) - 1
            
        result_df = df.loc[start + 1: end]
        logger.info(f"Successfully sliced data for Lesson {lesson_num}")
        return result_df
        
    except Exception as e:
        logger.error(f"Error slicing data for Lesson {lesson_num}: {e}")
        raise


def extract_video_data(lesson_df: pd.DataFrame) -> List[Tuple[str, str]]:
    """
    Extract video data (title and voiceover) from a lesson DataFrame.
    
    Args:
        lesson_df (pd.DataFrame): DataFrame containing lesson data
        
    Returns:
        List[Tuple[str, str]]: List of (title, voiceover) tuples
    """
    try:
        video_data = []
        video_rows = lesson_df[
            lesson_df[0].astype(str).str.strip().str.lower().str.startswith("video")
        ]
        
        if video_rows.empty:
            logger.warning("No video rows found in lesson data")
            return video_data
            
        for _, row in video_rows.iterrows():
            title = str(row[1]).strip() if pd.notna(row[1]) else "Untitled Video"
            voiceover = str(row[3]).strip() if pd.notna(row[3]) else ""
            
            # Skip if voiceover is empty or invalid
            if not voiceover or voiceover.lower() in {"nan", ""}:
                logger.warning(f"Skipping video '{title}' - empty voiceover script")
                continue
                
            video_data.append((title, voiceover))
            
        logger.info(f"Extracted {len(video_data)} videos from lesson data")
        return video_data
        
    except Exception as e:
        logger.error(f"Error extracting video data: {e}")
        return []


def create_excel_buffer(df: pd.DataFrame) -> BytesIO:
    """
    Create an in-memory Excel buffer from a DataFrame.
    
    Args:
        df (pd.DataFrame): DataFrame to convert to Excel
        
    Returns:
        BytesIO: Excel file buffer
    """
    try:
        buffer = BytesIO()
        with pd.ExcelWriter(buffer, engine="openpyxl") as writer:
            df.to_excel(writer, index=False, sheet_name="temp_sheet")
        buffer.seek(0)
        return buffer
        
    except Exception as e:
        logger.error(f"Error creating Excel buffer: {e}")
        raise


def create_docx_in_memory(html_text: str) -> BytesIO:
    """
    Create a DOCX file in memory from HTML content.
    
    Args:
        html_text (str): HTML content to convert
        
    Returns:
        BytesIO: DOCX file buffer
    """
    try:
        from bs4 import BeautifulSoup
        from docx import Document
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn
        
        soup = BeautifulSoup(html_text, "html.parser")
        doc = Document()
        
        for tbl_html in soup.find_all("table"):
            rows = tbl_html.find_all("tr")
            if not rows:
                continue
                
            ncols = len(rows[0].find_all(["th", "td"]))
            table = doc.add_table(rows=0, cols=ncols)
            
            for row in rows:
                cells = row.find_all(["th", "td"])
                row_cells = table.add_row().cells
                for i, cell in enumerate(cells):
                    if i < len(row_cells):
                        row_cells[i].text = cell.get_text(strip=True)
                        
            # Add borders
            tbl = table._tbl
            tblPr = tbl.tblPr
            borders = OxmlElement("w:tblBorders")
            
            for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
                b = OxmlElement(f"w:{edge}")
                b.set(qn("w:val"), "single")
                b.set(qn("w:sz"), "4")
                b.set(qn("w:color"), "000000")
                borders.append(b)
                
            tblPr.append(borders)
        
        # Save to BytesIO buffer
        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer
        
    except Exception as e:
        logger.error(f"Error creating DOCX in memory: {e}")
        raise


def slugify(text: str, max_len: int = 60) -> str:
    """
    Convert text to a slug-friendly format.

    Args:
        text (str): The input text
        max_len (int): The maximum length of the slug

    Returns:
        str: The slugified text
    """
    try:
        if not text or pd.isna(text):
            return "untitled"
            
        text = re.sub(r"[^\w\s-]", "", str(text)).strip().lower()
        slug = re.sub(r"[-\s]+", "_", text)[:max_len]
        return slug if slug else "untitled"
        
    except Exception as e:
        logger.error(f"Error creating slug from text: {e}")
        return "untitled"


def read_excel_from_s3(url: str, sheet_name: str) -> pd.DataFrame:
    """
    Read an Excel file from S3 and return a DataFrame.
    
    Args:
        url (str): URL of the Excel file
        sheet_name (str): Name of the sheet to read
        
    Returns:
        pd.DataFrame: DataFrame containing the Excel data
    """
    try:
        logger.info(f"Downloading the file from {url}")
        # Download the file
        response = requests.get(url)
        response.raise_for_status()  # Raise an exception for bad status codes
        
        # Create a temporary file to store the downloaded content
        with tempfile.NamedTemporaryFile(suffix='.xlsx', delete=False) as temp_file:
            temp_file.write(response.content)
            temp_file_path = temp_file.name
            
        # Read the Excel file into a DataFrame
        df = pd.read_excel(temp_file_path, sheet_name=sheet_name, header=None)
        
        # Clean up the temporary file
        os.unlink(temp_file_path)
        
        return df
        
    except requests.exceptions.RequestException as e:
        logger.error(f"Error downloading the file: {e}")
        return None
    except Exception as e:
        logger.error(f"Error reading the file: {e}")
        return None